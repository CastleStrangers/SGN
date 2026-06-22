import sys
import subprocess
import os

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def create_document():
    print("=====================================================")
    print("جاري إنشاء دليل اختصارات SGN على سطح المكتب...")
    print("=====================================================")
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title.add_run("دليل اختصارات سطح المكتب لمشروع SGN")
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 50) # HSL-like green #1a5632
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sub = subtitle.add_run("سهولة التحكم والتشغيل والرفع بضغطة زر واحدة")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT # spacer
    
    # Intro
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("هذا الدليل يشرح عمل الاختصارات الثلاثة التي تم إنشاؤها وتثبيتها على سطح المكتب لتسهيل وسرعة العمل اليومي على مشروع الجالية السورية في هولندا (SGN).")
    r.font.name = 'Arial'
    r.font.size = Pt(11.5)
    
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT # spacer
    
    # Table of Shortcuts
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["الوظيفة الأساسية", "اسم الاختصار", "رقم"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1a5632")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
    shortcuts_info = [
        ("تشغيل خادم الويب (Next.js)، خادم الموبايل (Expo)، والمزامنة تلقائياً.", "SGN Start Dev", "1"),
        ("رفع التعديلات تفاعلياً لـ GitHub ونشر التحديث مباشرة على Vercel.", "SGN Push & Deploy", "2"),
        ("مزامنة وإعداد قاعدة البيانات السحابية (Turso) محلياً.", "SGN Sync DB (Turso)", "3")
    ]
    
    for desc, name, num in shortcuts_info:
        row_cells = table.add_row().cells
        row_cells[0].text = desc
        row_cells[1].text = name
        row_cells[2].text = num
        
        # Style cells
        for idx, cell in enumerate(row_cells):
            p = cell.paragraphs[0]
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True
            elif idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT # spacer
    
    # Detailed Sections
    shortcuts_details = [
        {
            "num": "1",
            "title": "الاختصار الأول: SGN Start Dev",
            "icon": "💻 تشغيل خادم التطوير الشامل",
            "desc": "يقوم هذا الاختصار بتشغيل بيئة التطوير الكاملة للمشروع بضغطة زر واحدة ودون الحاجة لفتح عدة نوافذ ترمينال يدوياً.",
            "actions": [
                "تشغيل خادم الويب (Next.js) في وضع التطوير (Port 3000).",
                "تشغيل خادم تطبيق الهواتف الذكية (Expo React Native) (Port 3001).",
                "بدء تشغيل عملية المزامنة التلقائية مع فيسبوك (PM2 Sync) في الخلفية.",
                "تشغيل قاعدة البيانات المحلية SQLite للربط المباشر."
            ],
            "when": "عندما تريد البدء بتعديل الأكواد، إضافة ميزات جديدة، أو تجربة التطبيق محلياً (سواء الويب أو الموبايل)."
        },
        {
            "num": "2",
            "title": "الاختصار الثاني: SGN Push & Deploy",
            "icon": "🚀 الرفع والنشر المباشر",
            "desc": "هذا الاختصار تفاعلي بالكامل، ومهمته الأساسية هي حفظ ورفع عملك وتحديث رابط العميل المباشر بعد كل تعديل.",
            "actions": [
                "فحص الملفات المعدلة وتجهيزها للحفظ (git add -A).",
                "سؤالك تفاعلياً عن 'رسالة الالتزام' (Commit Message) لتوضيح ما قمت بتعديله.",
                "سحب أي تحديثات جديدة من المستودع لتفادي التعارضات (git pull --rebase).",
                "رفع التعديلات فورياً إلى GitHub (git push).",
                "بناء ونشر المشروع تلقائياً على Vercel (vercel --prod) لتحديث الرابط المباشر للزبون."
            ],
            "when": "تستخدمه فور انتهائك من تحديث كود، أو إضافة ميزة، أو حل مشكلة في المشروع لرفعها وتحديث الرابط فوراً."
        },
        {
            "num": "3",
            "title": "الاختصار الثالث: SGN Sync DB (Turso)",
            "icon": "🗄️ مزامنة قاعدة البيانات والبيئة",
            "desc": "يقوم هذا الاختصار بربط ومزامنة البيانات محلياً مع قاعدة البيانات السحابية (Turso) وجلب إعدادات Vercel الأخيرة.",
            "actions": [
                "تنزيل وتحديث إعدادات البيئة (Environment Variables) الخاصة بالمشروع من Vercel.",
                "مزامنة هيكل الجداول الحالي وقاعدة البيانات السحابية Turso محلياً باستخدام Prisma db push."
            ],
            "when": "تستخدمه عند تحديث جداول قاعدة البيانات (Prisma Schema)، أو عند تنزيل المشروع على جهاز جديد لأول مرة لتهيئة قاعدة البيانات."
        }
    ]
    
    for item in shortcuts_details:
        # Add heading
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_title = h.add_run(f"\n{item['title']}")
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(16)
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(26, 86, 50)
        
        # Icon/Label
        p_label = doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_icon = p_label.add_run(item['icon'])
        r_icon.font.name = 'Arial'
        r_icon.font.size = Pt(12)
        r_icon.bold = True
        r_icon.font.color.rgb = RGBColor(0, 102, 204)
        
        # Desc
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_desc = p_desc.add_run(item['desc'])
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(11)
        
        # Actions Title
        p_act_title = doc.add_paragraph()
        p_act_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_act_title = p_act_title.add_run("ماذا يحدث عند تشغيله:")
        r_act_title.font.name = 'Arial'
        r_act_title.font.size = Pt(11)
        r_act_title.bold = True
        
        # Actions List
        for act in item['actions']:
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_bullet = p_bullet.add_run(act)
            r_bullet.font.name = 'Arial'
            r_bullet.font.size = Pt(10.5)
            
        # When to use
        p_when = doc.add_paragraph()
        p_when.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_when_lbl = p_when.add_run("متى يجب تشغيله؟ ")
        r_when_lbl.font.name = 'Arial'
        r_when_lbl.font.size = Pt(11)
        r_when_lbl.bold = True
        r_when_lbl.font.color.rgb = RGBColor(204, 102, 0)
        
        r_when_desc = p_when.add_run(item['when'])
        r_when_desc.font.name = 'Arial'
        r_when_desc.font.size = Pt(10.5)
        
    doc_path = os.path.join(os.path.expanduser("~"), "Desktop", "دليل_اختصارات_SGN.docx")
    doc.save(doc_path)
    print(f"Document created successfully on Desktop: {doc_path}")

if __name__ == "__main__":
    create_document()
