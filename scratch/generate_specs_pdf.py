import sys
import subprocess
import os

# Ensure python-docx and pywin32 are installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    import win32com.client
except ImportError:
    print("Installing pywin32...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pywin32"])
    import win32com.client

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def create_specification_document():
    print("=====================================================")
    print("جاري إنشاء وثيقة مواصفات مشروع SGN...")
    print("=====================================================")
    
    doc = Document()
    
    # Page setup (1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_title.add_run("وثيقة مواصفات وتحديثات مشروع SGN")
    run.font.name = 'Arial'
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 86, 50) # SGN Green
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sub = p_sub.add_run("الدليل الشامل لميزات المنصة، الهيكلية، وأحدث الإضافات")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT # separator
    
    # Intro
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_intro = p_intro.add_run("منصة SGN (الجالية السورية في هولندا) هي منصة رقمية متكاملة تهدف لخدمة وتيسير شؤون الجالية، وربط الأعضاء، ومشاركة الأخبار والفعاليات والفرص التطوعية. تتكون المنصة من تطبيق ويب متطور (Next.js 16) وتطبيق هواتف ذكية (Expo React Native).")
    r_intro.font.name = 'Arial'
    r_intro.font.size = Pt(11.5)
    
    # Section 1: Project Structure
    add_heading(doc, "1. هيكلية المشروع (Project Structure)")
    
    p_struct = doc.add_paragraph()
    p_struct.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_struct = p_struct.add_run("ينقسم الكود المصدري للمشروع إلى ثلاثة أجزاء رئيسية متكاملة:")
    r_struct.font.name = 'Arial'
    r_struct.font.size = Pt(11)
    
    add_bullet(doc, "تطبيق الويب (Web App - src/): مبني باستخدام Next.js 16 و TailwindCSS ويحتوي على لوحة التحكم والموقع التعريفي العام.")
    add_bullet(doc, "تطبيق الموبايل (Mobile App - mobile/): تطبيق تفاعلي للأعضاء مبني باستخدام Expo React Native ويحتوي على شاشات تصفح الأخبار والمحادثات المباشرة.")
    add_bullet(doc, "قاعدة البيانات المشتركة (prisma/): ملف schema.prisma المشترك الذي يحدد جداول قاعدة البيانات (SQLite محلياً و Turso/libsql للإنتاج).")
    
    # Section 2: Key Features
    add_heading(doc, "2. الميزات والخدمات الذكية (Key Features)")
    
    add_bullet(doc, "إدارة العضويات والبطاقات الرقمية: نموذج انتساب ذكي يتيح للأعضاء التسجيل والحصول على بطاقة عضوية رقمية تدعم التحقق الفوري عبر رمز الـ QR Code.")
    add_bullet(doc, "المخزن الآمن وتحليل المستندات (Vault): ميزة تتيح للأعضاء رفع وحفظ صور هوياتهم وخطاباتهم الرسمية بشكل آمن، مع تحليل فوري باللغة العربية يشرح محتوى الخطاب والإجراءات المطلوبة (بواسطة Ollama llava:7b).")
    add_bullet(doc, "المساعد الذكي المتخصص (AI Assistant): مساعد مدعوم بـ Ollama مع شخصيات متعددة (مساعد قانوني، مرشد لغة، مستشار توظيف وسكن)، مدعوم بنظام RAG للبحث في أخبار وفعاليات المنصة.")
    add_bullet(doc, "نظام التطوع الذكي (AI-Skill Matcher): نظام يقوم بتحليل مهارات وخبرات العضو واقتراح المهام التطوعية المناسبة له بدقة.")
    add_bullet(doc, "دليل الخدمات المهنية (Services Directory): ميزة جديدة تتيح للأعضاء (المقاولين، الحرفيين، المترجمين، إلخ) عرض خدماتهم المهنية وتسهيل الوصول إليهم.")
    add_bullet(doc, "المحادثة المباشرة (Direct Chat): نظام محادثات فورية مباشر بين الأعضاء ومزودي الخدمات يدعم الإشعارات الفورية (Push Notifications) عبر Expo.")
    add_bullet(doc, "نظام المزامنة التلقائية (Sync System): سحب المنشورات تلقائياً وتصنيفها من صفحة الفيسبوك الرسمية، قناة يوتيوب، وموقع sy-nl.org.")
    
    # Section 3: Board of Directors
    add_heading(doc, "3. مجلس الإدارة (Board of Directors)")
    p_board = doc.add_paragraph()
    p_board.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_board = p_board.add_run("يتكون مجلس إدارة الجالية حالياً من 24 عضواً مسجلين بالكامل في قاعدة البيانات وموزعين على اللجان والمكاتب المختلفة:")
    r_board.font.name = 'Arial'
    r_board.font.size = Pt(11)
    
    board_members = [
        ("عبد المنعم الشامان", "رئيس مجلس الإدارة"),
        ("صالح المحمد حنايا", "نائب رئيس مجلس الإدارة"),
        ("خالد فيصل الطويل", "الأمين العام للجالية"),
        ("محمد سليم عزيزة", "عضو مكتب الأمانة العامة (المسؤول التقني) - تم إضافته مؤخراً"),
        ("محمد رائد كعكة", "رئيس المكتب المالي"),
        ("أحمد الحرفي", "عضو مكتب الأمانة العامة (المسؤول التنظيمي)"),
        ("هدى الحلاق", "مديرة المكتب الإعلامي"),
        ("رائد دهموش المشهور", "مسؤول العلاقات الخارجية"),
        ("محمود الناصر", "مسؤول العلاقات الدولية"),
        ("عمر النويلاتي", "مسؤول النشاطات"),
        ("رابعة الزريقات", "مسؤولة الشؤون القانونية"),
        ("فاتن رحال", "رئيسة مكتب شؤون المرأة والأسرة"),
        ("حسن الحسن قطيني", "مسؤول رواد الأعمال"),
        ("نبيل حاج حسين", "مسؤول الصحة والدعم النفسي"),
        ("بلال الرفاعي", "مسؤول الصحة والدعم النفسي"),
        ("محمد مصطفى سمهاني", "مسؤول العلاقات الدولية"),
        ("نهاد سويد", "مسؤول العلاقات العامة"),
        ("محمد أكرم الجنيدي", "مسؤول الطلاب والتعليم"),
        ("يوسف درويش", "لجنة الرقابة الداخلية"),
        ("محمد ربيع الجنيدي", "المجلس الاستشاري"),
        ("وسيم حسان", "المجلس الاستشاري"),
        ("ماهره الطواشي", "المجلس الاستشاري"),
        ("ريمه الحربات", "المجلس الاستشاري"),
        ("فراس هاني عابدين", "المجلس الاستشاري")
    ]
    
    # Add a styled table for board members
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["المنصب / اللّجنة", "الاسم الكامل", "رقم"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1a5632")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
    for idx, (name, role) in enumerate(board_members, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = name
        row_cells[2].text = str(idx)
        
        # Formatting cells
        for cell_idx, cell in enumerate(row_cells):
            p = cell.paragraphs[0]
            if cell_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.runs[0].font.bold = True
            elif cell_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Section 4: Latest Additions
    add_heading(doc, "4. آخر الإضافات والتحديثات للمشروع (Latest Updates)")
    
    add_bullet(doc, "إضافة العضو التقني: تم تسجيل وإضافة العضو الـ 24 السيد محمد سليم عزيزة كمسؤول تقني في مجلس الإدارة وتحديث واجهات الموقع وقاعدة البيانات.")
    add_bullet(doc, "تحديث محرك البحث الذكي (RAG): تفعيل جلب المعلومات المحدثة من قواعد البيانات للأخبار، الفعاليات، الأعضاء، ودليل الخدمات ودمجها في إجابات الذكاء الاصطناعي.")
    add_bullet(doc, "نظام المزامنة الذاتية والتشخيص: إضافة سكريبتات متكاملة لفحص توكن الفيسبوك وتشخيص الأخطاء تلقائياً لتفادي انقطاع مزامنة المقالات.")
    add_bullet(doc, "تصحيح وعرض صور أعضاء مجلس الإدارة بمسارات وتنسيقات SVG تفاعلية وحديثة.")
    
    # Section 5: Development Guidelines
    add_heading(doc, "5. إرشادات وقواعد التطوير والرفع (Rules)")
    
    add_bullet(doc, "قاعدة المزامنة الشاملة: أي ميزة أو تعديل أو إضافة للمفاتيح اللغوية (i18n) يجب أن تتم في الويب والموبايل معاً.")
    add_bullet(doc, "قاعدة النشر الفوري: يجب رفع كل تعديل فوراً لـ GitHub وتحديث خادم الإنتاج ورابط العميل Vercel (sgn-indol.vercel.app).")
    
    # Save word doc
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders")
        desktop_raw, _ = winreg.QueryValueEx(key, "Desktop")
        desktop_path = os.path.expandvars(desktop_raw)
    except Exception:
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        
    docx_filename = "SGN_Project_Specifications.docx"
    pdf_filename = "SGN_Project_Specifications.pdf"
    
    desktop_docx_path = os.path.join(desktop_path, docx_filename)
    desktop_pdf_path = os.path.join(desktop_path, pdf_filename)
    
    public_docx_path = os.path.join("public", "pdfs", docx_filename)
    public_pdf_path = os.path.join("public", "pdfs", pdf_filename)
    
    # Save DOCX locally first
    doc.save(desktop_docx_path)
    print(f"Saved DOCX to Desktop: {desktop_docx_path}")
    
    # Save to public directory
    os.makedirs(os.path.join("public", "pdfs"), exist_ok=True)
    doc.save(public_docx_path)
    print(f"Saved DOCX to Public folder: {public_docx_path}")
    
    # Convert DOCX to PDF using Word automation
    try:
        print("Converting DOCX to PDF via MS Word...")
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Convert Desktop file
        word_doc = word.Documents.Open(os.path.abspath(desktop_docx_path))
        word_doc.SaveAs(os.path.abspath(desktop_pdf_path), FileFormat=17) # 17 = wdFormatPDF
        word_doc.Close()
        
        # Convert Public file
        word_doc = word.Documents.Open(os.path.abspath(public_docx_path))
        word_doc.SaveAs(os.path.abspath(public_pdf_path), FileFormat=17)
        word_doc.Close()
        
        word.Quit()
        print(f"Successfully created PDF on Desktop: {desktop_pdf_path}")
        print(f"Successfully created PDF in Public folder: {public_pdf_path}")
        
        # Create Desktop shortcut to the project PDF file
        try:
            print("Creating Desktop shortcut to the project PDF...")
            shell = win32com.client.Dispatch("WScript.Shell")
            shortcut = shell.CreateShortCut(os.path.join(desktop_path, "SGN Project Specifications.lnk"))
            shortcut.TargetPath = os.path.abspath(public_pdf_path)
            shortcut.WorkingDirectory = os.path.abspath(os.path.join("public", "pdfs"))
            shortcut.Description = "SGN Project Specifications PDF"
            shortcut.save()
            print("Desktop shortcut created successfully!")
        except Exception as shortcut_error:
            print(f"Could not create shortcut: {shortcut_error}")
            
    except Exception as e:
        print(f"Error converting to PDF via MS Word: {e}")
        print("Please note that DOCX has been saved successfully.")

def add_heading(doc, text):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = h.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(15)
    r.bold = True
    r.font.color.rgb = RGBColor(26, 86, 50)
    
def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)

if __name__ == "__main__":
    create_specification_document()
